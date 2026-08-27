import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=130, bottom=130, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="E2E8F0", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def create_accelerated_simplified_doc():
    doc = Document()
    
    # Margens da página
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    PRIMARY_COLOR = RGBColor(234, 88, 12)     # Laranja #ea580c
    SECONDARY_COLOR = RGBColor(37, 99, 235)   # Azul #2563eb
    DARK_TEXT = RGBColor(30, 41, 59)          # Grafite escuro #1e293b
    MUTED_TEXT = RGBColor(100, 116, 139)      # Cinza suave #64748b

    # ==================== CABEÇALHO ====================
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(2)
    
    run_pre = header_p.add_run("CRONOGRAMA EXECUTIVO ACELERADO & ACOMPANHAMENTO DE ENTREGAS\n")
    run_pre.font.name = "Arial"
    run_pre.font.size = Pt(11)
    run_pre.font.bold = True
    run_pre.font.color.rgb = SECONDARY_COLOR

    run_title = header_p.add_run("Sistema de Gestão do Centro de Distribuição & 5 Lojas")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(17)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    run_sub = sub_p.add_run("Planejamento de Alta Prioridade — Entrega Final Impreterível até 10 de Setembro de 2026")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(10)
    run_sub.font.italic = True
    run_sub.font.color.rgb = MUTED_TEXT

    # ==================== CARD DE RESUMO ====================
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    col_widths = [Inches(3.4), Inches(3.4)]
    
    cells_data = [
        [("Data do Relatório:", " 18 de Agosto de 2026"), ("Data Limite de Entrega Final:", " 10 de Setembro de 2026")],
        [("Ritmo de Execução:", " Sprints Curtas Aceleradas (3 a 4 dias)"), ("Fase Atual:", " Melhorias nos Cadastros & Ajuste do Envio p/ Lojas")]
    ]
    
    for row_idx, row in enumerate(info_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = col_widths[col_idx]
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=90, bottom=90, left=130, right=130)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            lbl, val = cells_data[row_idx][col_idx]
            r1 = p.add_run(lbl)
            r1.font.name = "Arial"
            r1.font.size = Pt(9.5)
            r1.font.bold = True
            r1.font.color.rgb = DARK_TEXT
            
            r2 = p.add_run(val)
            r2.font.name = "Arial"
            r2.font.size = Pt(9.5)
            if "10 de Setembro" in val:
                r2.font.bold = True
                r2.font.color.rgb = PRIMARY_COLOR
            elif "Melhorias" in val:
                r2.font.bold = True
                r2.font.color.rgb = SECONDARY_COLOR
            else:
                r2.font.color.rgb = DARK_TEXT

    set_table_borders(info_table, color="E2E8F0", sz="6", val="single")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==================== SEÇÃO 1: SITUAÇÃO ATUAL ====================
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    rh1 = h1.add_run("1. Em que fase o projeto está hoje?")
    rh1.font.name = "Arial"
    rh1.font.size = Pt(12.5)
    rh1.font.bold = True
    rh1.font.color.rgb = PRIMARY_COLOR

    p_diag = doc.add_paragraph()
    p_diag.paragraph_format.space_after = Pt(6)
    r_diag = p_diag.add_run(
        "Para garantir o cumprimento do prazo final em 10/09/2026, a equipe reduziu o tempo de cada etapa para ciclos "
        "rápidos e focados de 3 a 4 dias.\n\n"
        "Hoje (18/08), estamos aplicando melhorias contínuas nas telas de cadastro de produtos (grades, tamanhos e preços) "
        "e concluindo a estabilização da comunicação com as 5 lojas (garantindo que qualquer cadastro novo ou envio chegue "
        "às lojas sem falhas). Com isso, liberamos a validação de compras e romaneios ainda neste mês de agosto."
    )
    r_diag.font.name = "Arial"
    r_diag.font.size = Pt(9.5)
    r_diag.font.color.rgb = DARK_TEXT

    # ==================== SEÇÃO 2: TABELA DE ENTREGAS ACELERADA ====================
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    rh2 = h2.add_run("2. Cronograma Acelerado de Entregas (Prazo Máximo: 10/09/2026)")
    rh2.font.name = "Arial"
    rh2.font.size = Pt(12.5)
    rh2.font.bold = True
    rh2.font.color.rgb = PRIMARY_COLOR

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="CBD5E1", sz="4", val="single")

    headers = ["Etapa do Sistema", "O que faz / Benefício", "Previsão de Entrega", "Situação"]
    widths = [Inches(1.6), Inches(2.7), Inches(1.1), Inches(1.4)]
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=90, right=90)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    stages_data = [
        (
            "1. Cadastros & Melhorias",
            "Cadastro centralizado de produtos, tamanhos, cores, fornecedores e melhorias de usabilidade no CD.",
            "20/08/2026",
            "PRONTO / EM AJUSTES",
            "10B981",  # Verde
            "F0FDF4"
        ),
        (
            "2. Envio de Dados CD ➔ Lojas",
            "Envio automático dos novos produtos para as 5 filiais e retorno das vendas diárias para a Matriz sem perdas.",
            "21/08/2026",
            "EM AJUSTE (FOCO HOJE)",
            "E11D48",  # Vermelho/Rose
            "FFF1F2"
        ),
        (
            "3. Entrada de Notas de Compra",
            "Leitura da Nota Fiscal do fornecedor (XML), cálculo de custos/fretes e geração do Contas a Pagar.",
            "26/08/2026",
            "EM VALIDAÇÃO",
            "EAB308",  # Amarelo
            "FEFCE8"
        ),
        (
            "4. Envio de Mercadorias (Romaneio)",
            "Montagem de caixas por loja, impressão de guia de separação e baixa automática do estoque no galpão.",
            "31/08/2026",
            "EM VALIDAÇÃO",
            "EAB308",  # Amarelo
            "FEFCE8"
        ),
        (
            "5. Nota Fiscal & Conferência Loja",
            "Emissão da Nota Fiscal de transporte e conferência física na chegada da loja (validação antes de liberar venda).",
            "05/09/2026",
            "PRÓXIMA FASE",
            "2563EB",  # Azul
            "EFF6FF"
        ),
        (
            "6. Painel de Vendas & Liberação Final",
            "Painel consolidado com gráficos de vendas de todas as lojas, sugestão de reposição e liberação 100% do sistema.",
            "10/09/2026",
            "ENTREGA FINAL",
            "10B981",  # Verde
            "F0FDF4"
        )
    ]

    for item in stages_data:
        stage_name, desc, delivery, status_text, status_color, bg_color = item
        row_cells = table.add_row().cells
        
        for i in range(4):
            row_cells[i].width = widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=80, bottom=80, left=80, right=80)
            
        # Etapa
        p0 = row_cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(stage_name)
        r0.font.name = "Arial"
        r0.font.size = Pt(8)
        r0.font.bold = True
        r0.font.color.rgb = DARK_TEXT
        
        # Descrição
        p1 = row_cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(desc)
        r1.font.name = "Arial"
        r1.font.size = Pt(8)
        r1.font.color.rgb = DARK_TEXT

        # Data Entrega
        p2 = row_cells[2].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(delivery)
        r2.font.name = "Arial"
        r2.font.size = Pt(8)
        r2.font.bold = True
        r2.font.color.rgb = DARK_TEXT

        # Status Badge
        p3 = row_cells[3].paragraphs[0]
        p3.paragraph_format.space_after = Pt(0)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(status_text)
        r3.font.name = "Arial"
        r3.font.size = Pt(7.5)
        r3.font.bold = True
        r3.font.color.rgb = RGBColor(int(status_color[0:2], 16), int(status_color[2:4], 16), int(status_color[4:6], 16))

    # ==================== SEÇÃO 3: O QUE VALIDAR HOJE ====================
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    rh3 = h3.add_run("3. O que já está liberado para testes pelo cliente")
    rh3.font.name = "Arial"
    rh3.font.size = Pt(12.5)
    rh3.font.bold = True
    rh3.font.color.rgb = PRIMARY_COLOR

    points = [
        ("Cadastro de Produtos & Melhorias:", " Testar a inclusão de produtos com grade de tamanhos, cores, fotos e múltiplos de embalagem."),
        ("Entrada de Compras por Nota Fiscal:", " Testar a importação do XML da NF-e de fornecedor, cálculo automático de custos e títulos no Contas a Pagar."),
        ("Montagem de Romaneios:", " Testar a criação de remessas para as filiais (Itaporã, Maracaju, Nova Alvorada, Rio Brilhante, Douradina) e impressão da guia de separação.")
    ]

    for title, desc in points:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(2)
        r_bt = bp.add_run(title)
        r_bt.font.name = "Arial"
        r_bt.font.size = Pt(9)
        r_bt.font.bold = True
        r_bt.font.color.rgb = DARK_TEXT
        
        r_bd = bp.add_run(desc)
        r_bd.font.name = "Arial"
        r_bd.font.size = Pt(9)
        r_bd.font.color.rgb = DARK_TEXT

    # ==================== SEÇÃO 4: PRÓXIMOS PASSOS & ASSINATURA ====================
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)
    rh4 = h4.add_run("4. Próxima Entrega Agendada: 21 de Agosto de 2026")
    rh4.font.name = "Arial"
    rh4.font.size = Pt(12.5)
    rh4.font.bold = True
    rh4.font.color.rgb = PRIMARY_COLOR

    p_next = doc.add_paragraph()
    p_next.paragraph_format.space_after = Pt(14)
    r_next = p_next.add_run(
        "Nesta sexta-feira (21/08), concluiremos a estabilização da comunicação com as lojas e os refinamentos de cadastro. "
        "Na semana seguinte validaremos compras e romaneios de ponta a ponta, garantindo que em 10/09 o sistema completo "
        "esteja 100% entregue e operando em produção."
    )
    r_next.font.name = "Arial"
    r_next.font.size = Pt(9.5)
    r_next.font.color.rgb = DARK_TEXT

    # Bloco de Assinaturas
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    for cell in sign_table.rows[0].cells:
        cell.width = Inches(3.4)
        set_cell_margins(cell, top=160, bottom=40, left=40, right=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r_line = p.add_run("_________________________________________\n")
        r_line.font.color.rgb = MUTED_TEXT
        
    p_l = sign_table.rows[0].cells[0].paragraphs[0]
    r_l_title = p_l.add_run("Equipe do Projeto")
    r_l_title.font.name = "Arial"
    r_l_title.font.size = Pt(8.5)
    r_l_title.font.bold = True

    p_r = sign_table.rows[0].cells[1].paragraphs[0]
    r_r_title = p_r.add_run("Cliente / Aprovador")
    r_r_title.font.name = "Arial"
    r_r_title.font.size = Pt(8.5)
    r_r_title.font.bold = True

    output_path = r"g:\PROJETOS\CENTRO-DISTRIBUICAO\docs\CRONOGRAMA_ENTREGA_CENTRO_DISTRIBUICAO.docx"
    doc.save(output_path)
    print(f"Documento acelerado salvo com sucesso em: {output_path}")

if __name__ == "__main__":
    create_accelerated_simplified_doc()
